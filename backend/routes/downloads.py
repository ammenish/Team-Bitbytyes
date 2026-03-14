"""
Routes for generating and downloading PDF/DOCX files for MoM and Gist documents.
"""
import io
import os
import textwrap
from flask import Blueprint, send_file, jsonify
from flask_jwt_extended import jwt_required, get_jwt_identity
from models.user import User
from models.application import Application
from fpdf import FPDF
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

downloads_bp = Blueprint("downloads", __name__, url_prefix="/api/download")


class PariveshPDF(FPDF):
    """Custom PDF with official MoEF&CC header, footer, and watermark."""

    def __init__(self, title="PARIVESH 3.0", *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.doc_title = title
        # Load reliable built-in fallback fonts
        self.set_margins(left=20, top=20, right=20)

    def header(self):
        # Only print the full formal header on the first page
        if self.page_no() == 1:
            try:
                # Attempt to place logos (assuming ashoka.png and moefcc_logo.png exist in backend/static/)
                self.image("static/ashoka.png", x=20, y=10, w=20)
                self.image("static/moefcc_logo.png", x=160, y=10, w=25)
            except Exception:
                pass # If images missing, proceed without crashing
            
            # Central Text
            self.set_y(15)
            self.set_font("Helvetica", "B", 12)
            self.set_text_color(0, 0, 0)
            self.cell(0, 5, _safe("Government of India"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.cell(0, 5, _safe("Ministry of Environment, Forest and Climate Change"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_font("Helvetica", "B", 10)
            self.cell(0, 5, _safe("IA Division"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.cell(0, 5, _safe("(Non-Coal Mining)"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_font("Helvetica", "B", 12)
            self.cell(0, 5, "***", align="C", new_x="LMARGIN", new_y="NEXT")
            self.ln(10)

        # Background Watermark
        try:
            with self.local_context(fill_opacity=0.08):
                # Put a faded watermark in the center of every page
                self.image("static/moefcc_logo.png", x=60, y=100, w=90)
        except Exception:
            pass

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "", 9)
        self.set_text_color(100, 100, 100)
        # Right aligned "Page X of Y"
        self.cell(0, 10, f"Page {self.page_no()} of {{nb}}", align="R")


def _safe(text):
    """Strip non-latin-1 characters for fpdf2 built-in fonts."""
    return text.encode("latin-1", errors="replace").decode("latin-1") if text else ""

def _build_pdf(app, doc_type="mom"):
    """Generate a PDF for MoM or Gist."""
    content = app.mom if doc_type == "mom" else app.gist
    is_mom = doc_type == "mom"
    
    # Precise titles mimicking official documents
    title_upper = ("Minutes of AGENDA FOR MEETING OF THE RE-CONSTITUTED EXPERT APPRAISAL " 
                   "COMMITTEE (NON-COAL MINING SECTOR)") if is_mom else "MEETING GIST / VERIFICATION REPORT"

    pdf = PariveshPDF(title="PARIVESH Minutes")
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    app_date = app.created_at.strftime('%d/%m/%Y') if hasattr(app, 'created_at') and app.created_at else '12/03/2026'

    # Date aligning to the right
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 5, _safe(f"Date: {app_date}"), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Meeting Title Block
    pdf.set_font("Helvetica", "B", 11)
    # The true PDF has the title centered and bolded
    pdf.multi_cell(0, 6, _safe(title_upper), align="C")
    pdf.ln(8)

    # Key Value pairs (MoM ID, etc)
    pdf.set_font("Helvetica", "B", 10)
    
    def print_kv(label, val):
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(40, 6, _safe(label))
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, _safe(val), new_x="LMARGIN", new_y="NEXT")
        
    print_kv("MoM ID:", f"EC/MOM/EAC/{app.app_id}/2026")
    print_kv("Agenda ID:", f"EC/AGENDA/EAC/{app.app_id}/2026")
    print_kv("Meeting Venue:", "Indira Paryavaran Bhawan, Jor Bagh, New Delhi")
    print_kv("Meeting Mode:", "Physical")
    pdf.ln(6)

    # Date & Time mini-table
    pdf.set_font("Helvetica", "", 10)
    pdf.set_fill_color(245, 245, 245)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.2)
    pdf.cell(55, 7, _safe(app_date), border=1, align="C")
    pdf.cell(55, 7, "10:00 AM", border=1, align="C")
    pdf.cell(60, 7, "06:30 PM", border=1, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)

    # Proposal details nested table header
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, _safe("2. The details of the project submitted by the Project Proponent are given under:"), new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, _safe("i. Project details:"), new_x="LMARGIN", new_y="NEXT")
    
    # Table layout: 2 columns
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(220, 230, 240) # Light blueish grey header shading
    
    # Save current Y for the row
    row_y = pdf.get_y()
    
    pdf.cell(50, 8, _safe("Name of the Proposal"), border=1, fill=True, new_x="RIGHT", new_y="TOP")
    pdf.set_font("Helvetica", "", 9)
    # Use standard cell if single line, but we must use get_string_width to check
    pdf.multi_cell(120, 8, _safe(app.project), border=1, new_x="LMARGIN", new_y="NEXT")
    
    # Calculate height of that multi_cell to know how far we advanced
    row2_y = pdf.get_y()
    
    # Align the X cursor back and print Location
    pdf.set_xy(20, row2_y)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(50, 8, _safe("Location"), border=1, fill=True, new_x="RIGHT", new_y="TOP")
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(120, 8, _safe(f"{app.sector} Sector / {app.category}"), border=1, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(10)

    # Content body
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 8, _safe("Salient Features:"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)

    if content:
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue

            # Detect section headings (numbered or all-caps lines)
            if (line and (line[0].isdigit() and "." in line[:4])) or line.isupper():
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(10, 36, 99)
                pdf.set_x(15)
                pdf.multi_cell(0, 6, _safe(line))
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(15, 23, 42)
            elif line.startswith(("•", "-", "*", "(")):
                pdf.set_x(20)
                pdf.multi_cell(0, 5, _safe(line))
            else:
                pdf.set_x(15)
                pdf.multi_cell(0, 5, _safe(line))
    else:
        pdf.set_text_color(148, 163, 184)
        pdf.cell(0, 10, f"No {doc_type.upper()} content available for this application.")

    # Generate bytes
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def _build_docx(app, doc_type="mom"):
    """Generate a DOCX for MoM or Gist."""
    content = app.mom if doc_type == "mom" else app.gist
    title = "Minutes of the Meeting" if doc_type == "mom" else "Meeting Gist"

    doc = DocxDocument()

    # Document Properties
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header section
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("PARI✓ESH 3.0\n")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(10, 36, 99)

    sub_run = header_para.add_run("Ministry of Environment, Forest & Climate Change")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    # Horizontal line
    doc.add_paragraph("_" * 80)

    # Title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Application details table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"

    cells = [
        ("Application ID", app.app_id),
        ("Project", app.project),
        ("Sector / Category", f"{app.sector} / {app.category}"),
    ]
    for i, (label, value) in enumerate(cells):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Content body
    if content:
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            if (line and (line[0].isdigit() and "." in line[:4])) or line.isupper():
                heading = doc.add_heading(line, level=2)
            elif line.startswith(("•", "-", "*")):
                doc.add_paragraph(line[1:].strip(), style="List Bullet")
            elif line.startswith("("):
                doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(line)
    else:
        p = doc.add_paragraph(f"No {doc_type.upper()} content available for this application.")
        p.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    # Footer
    doc.add_paragraph("_" * 80)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by PARIVESH 3.0 | Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@downloads_bp.route("/<int:app_id>/<doc_type>/<file_format>", methods=["GET"])
@jwt_required()
def download_document(app_id, doc_type, file_format):
    """
    Download a MoM or Gist as PDF or DOCX.
    URL: /api/download/<app_id>/<mom|gist>/<pdf|docx>
    Example: /api/download/1/mom/pdf
    """
    if doc_type not in ("mom", "gist"):
        return jsonify({"error": "doc_type must be 'mom' or 'gist'"}), 400
    if file_format not in ("pdf", "docx"):
        return jsonify({"error": "file_format must be 'pdf' or 'docx'"}), 400

    app = Application.query.get(app_id)
    if not app:
        return jsonify({"error": "Application not found"}), 404

    filename = f"{app.app_id}_{doc_type.upper()}.{file_format}"

    if file_format == "pdf":
        buf = _build_pdf(app, doc_type)
        mimetype = "application/pdf"
    else:
        buf = _build_docx(app, doc_type)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return send_file(buf, mimetype=mimetype, as_attachment=True, download_name=filename)
